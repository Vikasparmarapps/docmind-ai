# ============================================================
# export/docx_export.py  —  Generates styled DOCX from Q&A pairs
# ============================================================
# This file has ONE job: take a list of Q&A pairs → return DOCX bytes
#
# We use python-docx — a library for creating Microsoft Word files.
# Word documents are built paragraph by paragraph, with each run
# (piece of text) having its own font, size, and color settings.

import io
import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _add_divider(doc, color: str = "00e5ff", thickness: int = 12):
    """
    Add a horizontal divider line to the Word document.
    Word doesn't have a built-in HR element, so we fake it
    by adding a bottom border to an empty paragraph.
    """
    p = doc.add_paragraph()
    pPr   = p._p.get_or_add_pPr()
    pBdr  = OxmlElement("w:pBdr")
    bdr   = OxmlElement("w:bottom")
    bdr.set(qn("w:val"), "single")
    bdr.set(qn("w:sz"), str(thickness))
    bdr.set(qn("w:space"), "1")
    bdr.set(qn("w:color"), color)
    pBdr.append(bdr)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    return p


def export_docx(pairs: list, title: str = "DocMind AI — Q&A Export") -> bytes:
    """
    Generate a styled Word (.docx) file from Q&A pairs.

    Args:
        pairs : list of {"q": "question", "a": "answer"} dicts
        title : title shown at the top of the document

    Returns:
        DOCX file as bytes (ready for st.download_button)

    Visual structure of each Q&A entry:
        ┌────────────────────────────────┐
        │  Q01  (cyan, bold, small)      │
        │  The question text here        │  (white, bold, 11pt)
        │  A01  (green, bold, small)     │
        │  The answer text here          │  (grey, 10pt)
        │  ─────────────────────────     │
        └────────────────────────────────┘
    """
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run  = title_para.add_run(title)
    title_run.bold           = True
    title_run.font.size      = Pt(20)
    title_run.font.color.rgb = RGBColor(0x00, 0xe5, 0xff)

    # ── Subtitle (date + count)
    now = datetime.datetime.now().strftime("%B %d, %Y  %I:%M %p")
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run  = meta_para.add_run(f"Generated: {now}  ·  {len(pairs)} pairs")
    meta_run.font.size      = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    # ── Top divider (thick cyan)
    _add_divider(doc, color="00e5ff", thickness=12)

    # ── Each Q&A pair
    for i, pair in enumerate(pairs, 1):

        # Q label  e.g. "  Q01"
        q_label_para = doc.add_paragraph()
        q_label_run  = q_label_para.add_run(f"  Q{i:02d}")
        q_label_run.bold           = True
        q_label_run.font.size      = Pt(8)
        q_label_run.font.color.rgb = RGBColor(0x00, 0xe5, 0xff)
        q_label_para.paragraph_format.space_before = Pt(10)
        q_label_para.paragraph_format.space_after  = Pt(2)

        # Question text
        q_para = doc.add_paragraph()
        q_para.paragraph_format.left_indent  = Cm(0.5)
        q_para.paragraph_format.space_after  = Pt(6)
        q_run  = q_para.add_run(pair["q"])
        q_run.bold           = True
        q_run.font.size      = Pt(11)
        q_run.font.color.rgb = RGBColor(0xf1, 0xf5, 0xf9)

        # A label  e.g. "  A01"
        a_label_para = doc.add_paragraph()
        a_label_run  = a_label_para.add_run(f"  A{i:02d}")
        a_label_run.bold           = True
        a_label_run.font.size      = Pt(8)
        a_label_run.font.color.rgb = RGBColor(0x34, 0xd3, 0x99)
        a_label_para.paragraph_format.space_before = Pt(4)
        a_label_para.paragraph_format.space_after  = Pt(2)

        # Answer text
        a_para = doc.add_paragraph()
        a_para.paragraph_format.left_indent = Cm(0.5)
        a_para.paragraph_format.space_after = Pt(4)
        a_run  = a_para.add_run(pair["a"])
        a_run.font.size      = Pt(10)
        a_run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

        # Thin divider between pairs
        _add_divider(doc, color="1e293b", thickness=4)

    # ── Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run  = footer_para.add_run("DocMind AI · RAG · ChromaDB · Ollama · LangChain")
    footer_run.font.size      = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    # Save to bytes buffer and return
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
